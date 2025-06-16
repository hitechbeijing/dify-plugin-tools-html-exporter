import io
import logging
import re
from typing import Generator, Dict, Any, Optional, List, Tuple

from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from tools.html_to_docx.font_enum import DocxFontEnum
from tools.utils.file_utils import get_meta_data
from tools.utils.mimetype_utils import MimeType
from tools.utils.param_utils import get_html_text


class HtmlToDocxTool(Tool):
    logger = logging.getLogger(__name__)

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._processed_tags = set()
        self._base_font_size = Pt(12)  # 基础字号
        self._current_paragraph = None  # 跟踪当前段落
        self._current_style_state = None  # 当前样式状态
        self._list_level = 0  # 列表嵌套层级
        self._div_stack = []  # 用于跟踪div嵌套

    def _invoke(self, tool_parameters: dict) -> Generator[ToolInvokeMessage, None, None]:
        html_text = get_html_text(tool_parameters)
        result_file_bytes = html_text.encode("utf-8")
        try:
            doc = self.create_document_with_styles()
            self.html_to_docx(doc, result_file_bytes)

            result_bytes_io = io.BytesIO()
            doc.save(result_bytes_io)
            result_file_bytes = result_bytes_io.getvalue()
        except Exception as e:
            self.logger.exception("Failed to convert HTML to DOCX")
            yield self.create_text_message(f"Failed to convert HTML text to DOCX file, error: {str(e)}")
            return

        yield self.create_blob_message(
            blob=result_file_bytes,
            meta=get_meta_data(
                mime_type=MimeType.DOCX,
                output_filename=tool_parameters.get("output_filename"),
            ),
        )

    def create_document_with_styles(self):
        doc = Document()

        # 设置默认段落样式
        style = doc.styles['Normal']
        font = style.font
        font.name = DocxFontEnum.TIMES_NEW_ROMAN
        font.size = self._base_font_size
        # 移除默认颜色设置，避免覆盖
        rPr = style.element.get_or_add_rPr()
        rPr.rFonts.set(qn('w:eastAsia'), DocxFontEnum.SONG_TI)

        # 设置段落格式
        para_format = style.paragraph_format
        para_format.space_before = Pt(6)
        para_format.space_after = Pt(6)
        para_format.line_spacing = 1.5

        return doc

    def html_to_docx(self, doc: Document, html: str):
        soup = BeautifulSoup(html, 'html.parser')
        body = soup.find('body') or soup
        
        # 重置处理状态
        self._processed_tags = set()
        self._current_paragraph = None
        # 移除默认颜色设置
        self._current_style_state = {
            'bold': False,
            'italic': False,
            'underline': False,
            'font_size': None,
            'color': None,  # 不再设置默认颜色
            'font_family': None,
            'line_height': None,
            'highlight_color': None,  # 文字高亮颜色
            'background_color': None  # 块级元素背景色
        }
        self._list_level = 0
        self._div_stack = []
        
        # 处理所有子元素
        for child in body.children:
            if isinstance(child, Tag):
                self.process_tag(doc, child, self._current_style_state.copy())

    def process_tag(self, doc: Document, tag: Tag, parent_style: Dict[str, Any] = None):
        if self._is_processed(tag):
            return

        name = tag.name
        self._mark_as_processed(tag)

        # 继承父样式
        current_style = parent_style.copy() if parent_style else self._current_style_state.copy()
        
        # 处理div容器
        if name == 'div':
            # 保存当前状态
            self._div_stack.append({
                'current_paragraph': self._current_paragraph,
                'current_style': self._current_style_state.copy()
            })
            
            # 更新当前样式状态
            if tag.has_attr('style'):
                self._update_style_from_attributes(current_style, tag)
            
            # 设置当前样式状态为div的样式
            self._current_style_state = current_style.copy()
            
            # 处理div的所有子节点
            for child in tag.children:
                if isinstance(child, Tag):
                    self.process_tag(doc, child, current_style.copy())
                elif isinstance(child, NavigableString) and child.strip():
                    # 如果div内有直接文本，创建段落
                    if not self._current_paragraph:
                        self._current_paragraph = doc.add_paragraph()
                    run = self._current_paragraph.add_run(child.strip())
                    self._apply_run_style(run, current_style)
            
            # 恢复之前的样式状态
            if self._div_stack:
                prev_state = self._div_stack.pop()
                self._current_paragraph = prev_state['current_paragraph']
                self._current_style_state = prev_state['current_style']
            return

        # 处理标题标签
        if name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # 创建普通段落而不是标题
            paragraph = doc.add_paragraph()
            self._current_paragraph = paragraph
            
            # 创建标题样式状态
            level = int(name[1])
            heading_style = current_style.copy()
            heading_style['bold'] = True
            # 设置标题字号，级别越高字号越大
            base_sizes = {1: 24, 2: 20, 3: 18, 4: 16, 5: 14, 6: 12}
            heading_style['font_size'] = Pt(base_sizes.get(level, 12))
            
            # 应用标题特定样式
            if tag.has_attr('style'):
                self._update_style_from_attributes(heading_style, tag)
            
            # 处理标题内容
            self._process_children(doc, paragraph, tag, heading_style)
            self._current_paragraph = None
            
            # 设置标题对齐方式
            self.set_paragraph_alignment(paragraph, tag)
            
            # 应用块级背景色
            if heading_style.get('background_color'):
                self.set_paragraph_background(paragraph, heading_style['background_color'])
            return

        # 处理段落
        if name == 'p':
            # 创建新段落
            paragraph = doc.add_paragraph()
            self._current_paragraph = paragraph

            # 设置段落对齐方式
            self.set_paragraph_alignment(paragraph, tag)

            # 应用段落样式
            if tag.has_attr('style'):
                # 先更新当前样式状态
                self._update_style_from_attributes(current_style, tag)
                # 然后应用段落样式
                self.apply_paragraph_styles(paragraph, tag['style'])
            
            # 处理内容
            self._process_children(doc, paragraph, tag, current_style)
            self._current_paragraph = None
            
            # 应用块级背景色
            if current_style.get('background_color'):
                self.set_paragraph_background(paragraph, current_style['background_color'])
            return

        # 处理表格
        if name == 'table':
            self.handle_table(doc, tag, current_style)
            return

        # 处理列表
        if name in ['ul', 'ol']:
            self._list_level += 1
            self.handle_list(doc, tag, current_style)
            self._list_level -= 1
            return

        # 处理列表项
        if name == 'li':
            # 列表项由handle_list处理，这里跳过
            return

        # 处理换行
        if name == 'br':
            if self._current_paragraph:
                self._current_paragraph.add_run('\n')
            else:
                doc.add_paragraph('\n')
            return

        # 处理其他块级元素
        if name not in ['span', 'b', 'strong', 'i', 'em', 'u', 'small', 'font', 'mark']:
            text = tag.get_text().strip()
            if text:
                paragraph = doc.add_paragraph(text)
                self.set_paragraph_alignment(paragraph, tag)
                self._current_paragraph = None
                
                # 应用块级背景色
                if current_style.get('background_color'):
                    self.set_paragraph_background(paragraph, current_style['background_color'])

    def _process_children(self, doc: Document, paragraph: Paragraph, parent_tag: Tag, parent_style: Dict[str, Any]):
        """递归处理子节点，维护样式状态"""
        for child in parent_tag.children:
            if isinstance(child, Tag):
                # 复制父样式作为当前基础样式
                current_style = parent_style.copy()
                
                # 根据标签类型更新样式
                tag_name = child.name
                if tag_name in ['b', 'strong']:
                    current_style['bold'] = True
                elif tag_name in ['i', 'em']:
                    current_style['italic'] = True
                elif tag_name == 'u':
                    current_style['underline'] = True
                elif tag_name == 'small':
                    # 在基础字号上减少2pt
                    if current_style['font_size'] is None:
                        current_style['font_size'] = self._base_font_size - Pt(2)
                    else:
                        current_style['font_size'] = current_style['font_size'] - Pt(2)
                elif tag_name in ['span', 'font']:
                    # 处理span和font的内联样式 - 无论是否有style属性都处理
                    self._update_style_from_attributes(current_style, child)
                # 处理 <mark> 标签（高亮文本）
                elif tag_name == 'mark':
                    # 设置默认高亮颜色（黄色）
                    current_style['highlight_color'] = WD_COLOR_INDEX.YELLOW
                    
                    # 如果有内联样式，更新样式
                    if child.has_attr('style'):
                        self._update_style_from_attributes(current_style, child)
                    
                    # 递归处理子元素
                    self._process_children(doc, paragraph, child, current_style)
                    continue  # 跳过后续处理，因为已经递归处理了
                
                # 处理块级标签（如嵌套的p/div）
                if tag_name in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'table', 'ul', 'ol']:
                    # 创建新段落处理块级元素
                    self.process_tag(doc, child, current_style)
                else:
                    # 递归处理内联元素
                    self._process_children(doc, paragraph, child, current_style)
                    
            elif isinstance(child, NavigableString):
                # 文本节点，应用当前样式
                text = str(child).replace('\xa0', ' ')  # 替换不间断空格
                if text.strip() or text == '\n':  # 保留换行符
                    run = paragraph.add_run(text)
                    self._apply_run_style(run, parent_style)

    def _update_style_from_attributes(self, style_state: Dict[str, Any], tag: Tag):
        """从标签属性更新样式状态 - 优化颜色处理"""
        # 处理style属性
        if tag.has_attr('style'):
            styles = self._parse_style_string(tag['style'])
            
            # 字体颜色 - 优先处理
            if "color" in styles:
                color = styles["color"].strip()
                parsed_color = self._parse_color(color)
                if parsed_color:
                    style_state['color'] = parsed_color
            
            # 背景色（用于块级元素的背景）
            if "background-color" in styles:
                bg_color = styles["background-color"].strip()
                parsed_bg_color = self._parse_color(bg_color)
                if parsed_bg_color:
                    style_state['background_color'] = parsed_bg_color
            
            # 字体粗细
            if "font-weight" in styles:
                weight = styles["font-weight"].strip()
                style_state['bold'] = weight in ["bold", "bolder", "700", "800", "900"]
            
            # 字体样式
            if "font-style" in styles:
                font_style = styles["font-style"].strip()
                style_state['italic'] = font_style == "italic"
            
            # 文本装饰
            if "text-decoration" in styles:
                decoration = styles["text-decoration"].strip()
                style_state['underline'] = "underline" in decoration
            
            # 字体大小
            if "font-size" in styles:
                size = styles["font-size"].strip()
                if size.endswith("pt"):
                    try:
                        pt_size = float(size[:-2])
                        style_state['font_size'] = Pt(pt_size)
                    except ValueError:
                        pass
                elif size.endswith("px"):
                    try:
                        px_size = float(size[:-2])
                        # 简单转换：1pt ≈ 1.33px
                        style_state['font_size'] = Pt(px_size / 1.33)
                    except ValueError:
                        pass
                elif size.endswith("em"):
                    try:
                        em_size = float(size[:-2])
                        # 1em ≈ 12pt (基础字号)
                        style_state['font_size'] = self._base_font_size * em_size
                    except ValueError:
                        pass
            
            # 字体类型
            if "font-family" in styles:
                font_family = styles["font-family"].split(',')[0].strip('"\' ')
                style_state['font_family'] = font_family
            
            # 行高
            if "line-height" in styles:
                line_height = styles["line-height"].strip()
                try:
                    if line_height.replace('.', '', 1).isdigit():
                        style_state['line_height'] = float(line_height)
                except ValueError:
                    pass
        
        # 处理直接属性（如font标签的color、face、size）
        if tag.has_attr('color'):
            color = tag['color'].strip()
            parsed_color = self._parse_color(color)
            if parsed_color:
                style_state['color'] = parsed_color
        
        if tag.has_attr('face'):
            font_family = tag['face'].split(',')[0].strip('"\' ')
            style_state['font_family'] = font_family
        
        if tag.has_attr('size'):
            size = tag['size'].strip()
            # 映射HTML字体大小到实际字号
            size_map = {
                '1': Pt(8),
                '2': Pt(10),
                '3': Pt(12),
                '4': Pt(14),
                '5': Pt(18),
                '6': Pt(24),
                '7': Pt(36)
            }
            # 处理相对大小（如+2, -1）
            if size.startswith('+') or size.startswith('-'):
                try:
                    # 获取当前字号
                    current_size = style_state.get('font_size', self._base_font_size)
                    if not isinstance(current_size, Pt):
                        current_size = self._base_font_size
                    
                    # 计算相对变化
                    change = int(size)
                    # 每级变化约20%
                    new_size = current_size.pt * (1.0 + 0.2 * change)
                    style_state['font_size'] = Pt(new_size)
                except ValueError:
                    pass
            elif size in size_map:
                style_state['font_size'] = size_map[size]
            else:
                try:
                    # 尝试直接解析为pt值
                    pt_size = float(size)
                    style_state['font_size'] = Pt(pt_size)
                except ValueError:
                    pass

    def _map_color_to_highlight(self, color: RGBColor) -> WD_COLOR_INDEX:
        """将RGB颜色映射到Word的高亮颜色枚举值"""
        # 黄色高亮（<mark>的默认颜色）
        if color == RGBColor(255, 255, 0):  # 黄色
            return WD_COLOR_INDEX.YELLOW
        
        # 其他常见颜色的映射
        color_map = {
            (255, 0, 0): WD_COLOR_INDEX.RED,        # 红色
            (0, 255, 0): WD_COLOR_INDEX.BRIGHT_GREEN, # 绿色
            (0, 0, 255): WD_COLOR_INDEX.BLUE,        # 蓝色
            (255, 255, 0): WD_COLOR_INDEX.YELLOW,    # 黄色
            (255, 0, 255): WD_COLOR_INDEX.PINK,      # 粉色
            (0, 255, 255): WD_COLOR_INDEX.TURQUOISE, # 青绿色
            (255, 165, 0): WD_COLOR_INDEX.ORANGE,    # 橙色
        }
        
        # 查找最接近的颜色
        closest_color = None
        min_distance = float('inf')
        
        for rgb, highlight in color_map.items():
            r, g, b = rgb
            # 计算颜色距离（欧几里得距离）
            distance = ((color[0] - r) ** 2 + 
                        (color[1] - g) ** 2 + 
                        (color[2] - b) ** 2) ** 0.5
            
            if distance < min_distance:
                min_distance = distance
                closest_color = highlight
        
        return closest_color or WD_COLOR_INDEX.YELLOW  # 默认使用黄色

    def _parse_style_string(self, style_str: str) -> Dict[str, str]:
        """解析样式字符串为字典"""
        styles = {}
        for declaration in style_str.split(';'):
            declaration = declaration.strip()
            if not declaration:
                continue
            if ':' in declaration:
                key, value = declaration.split(':', 1)
                styles[key.strip().lower()] = value.strip()
        return styles

    def _parse_color(self, color_str: str) -> Optional[RGBColor]:
        """解析颜色值，支持多种格式 - 优化十六进制颜色解析"""
        if not color_str:
            return None
            
        color_str = color_str.strip().lower()
        
        # 支持颜色名称
        color_names = {
        'AliceBlue': RGBColor(240, 248, 255),
        'AntiqueWhite': RGBColor(250, 235, 215),
        'Aqua': RGBColor(0, 255, 255),
        'Aquamarine': RGBColor(127, 255, 212),
        'Azure': RGBColor(240, 255, 255),
        'Beige': RGBColor(245, 245, 220),
        'Bisque': RGBColor(255, 228, 196),
        'Black': RGBColor(0, 0, 0),
        'BlanchedAlmond': RGBColor(255, 235, 205),
        'Blue': RGBColor(0, 0, 255),
        'BlueViolet': RGBColor(138, 43, 226),
        'Brown': RGBColor(165, 42, 42),
        'BurlyWood': RGBColor(222, 184, 135),
        'CadetBlue': RGBColor(95, 158, 160),
        'Chartreuse': RGBColor(127, 255, 0),
        'Chocolate': RGBColor(210, 105, 30),
        'Coral': RGBColor(255, 127, 80),
        'CornflowerBlue': RGBColor(100, 149, 237),
        'Cornsilk': RGBColor(255, 248, 220),
        'Crimson': RGBColor(220, 20, 60),
        'Cyan': RGBColor(0, 255, 255),
        'DarkBlue': RGBColor(0, 0, 139),
        'DarkCyan': RGBColor(0, 139, 139),
        'DarkGoldenRod': RGBColor(184, 134, 11),
        'DarkGray': RGBColor(169, 169, 169),
        'DarkGreen': RGBColor(0, 100, 0),
        'DarkKhaki': RGBColor(189, 183, 107),
        'DarkMagena': RGBColor(139, 0, 139),
        'DarkOliveGreen': RGBColor(85, 107, 47),
        'DarkOrange': RGBColor(255, 140, 0),
        'DarkOrchid': RGBColor(153, 50, 204),
        'DarkRed': RGBColor(139, 0, 0),
        'DarkSalmon': RGBColor(233, 150, 122),
        'DarkSeaGreen': RGBColor(143, 188, 143),
        'DarkSlateBlue': RGBColor(72, 61, 139),
        'DarkSlateGray': RGBColor(47, 79, 79),
        'DarkTurquoise': RGBColor(0, 206, 209),
        'DarkViolet': RGBColor(148, 0, 211),
        'DeepPink': RGBColor(255, 20, 147),
        'DeepSkyBlue': RGBColor(0, 191, 255),
        'DimGray': RGBColor(105, 105, 105),
        'DodgerBlue': RGBColor(30, 144, 255),
        'FireBrick': RGBColor(178, 34, 34),
        'FloralWhite': RGBColor(255, 250, 240),
        'ForestGreen': RGBColor(34, 139, 34),
        'Fuchsia': RGBColor(255, 0, 255),
        'Gainsboro': RGBColor(220, 220, 220),
        'GhostWhite': RGBColor(248, 248, 255),
        'Gold': RGBColor(255, 215, 0),
        'GoldenRod': RGBColor(218, 165, 32),
        'Gray': RGBColor(128, 128, 128),
        'Green': RGBColor(0, 128, 0),
        'GreenYellow': RGBColor(173, 255, 47),
        'HoneyDew': RGBColor(240, 255, 240),
        'HotPink': RGBColor(255, 105, 180),
        'IndianRed': RGBColor(205, 92, 92),
        'Indigo': RGBColor(75, 0, 130),
        'Ivory': RGBColor(255, 255, 240),
        'Khaki': RGBColor(240, 230, 140),
        'Lavender': RGBColor(230, 230, 250),
        'LavenderBlush': RGBColor(255, 240, 245),
        'LawnGreen': RGBColor(124, 252, 0),
        'LemonChiffon': RGBColor(255, 250, 205),
        'LightBlue': RGBColor(173, 216, 230),
        'LightCoral': RGBColor(240, 128, 128),
        'LightCyan': RGBColor(224, 255, 255),
        'LightGoldenRodYellow': RGBColor(250, 250, 210),
        'LightGray': RGBColor(211, 211, 211),
        'LightGreen': RGBColor(144, 238, 144),
        'LightPink': RGBColor(255, 182, 193),
        'LightSalmon': RGBColor(255, 160, 122),
        'LightSeaGreen': RGBColor(32, 178, 170),
        'LightSkyBlue': RGBColor(135, 206, 250),
        'LightSlateGray': RGBColor(119, 136, 153),
        'LightSteelBlue': RGBColor(176, 196, 222),
        'LightYellow': RGBColor(255, 255, 224),
        'Lime': RGBColor(0, 255, 0),
        'LimeGreen': RGBColor(50, 205, 50),
        'Linen': RGBColor(250, 240, 230),
        'Magenta': RGBColor(255, 0, 255),
        'Maroon': RGBColor(128, 0, 0),
        'MediumAquaMarine': RGBColor(102, 205, 170),
        'MediumBlue': RGBColor(0, 0, 205),
        'MediumOrchid': RGBColor(186, 85, 211),
        'MediumPurple': RGBColor(147, 112, 219),
        'MediumSeaGreen': RGBColor(60, 179, 113),
        'MediumSlateBlue': RGBColor(123, 104, 238),
        'MediumSpringGreen': RGBColor(0, 250, 154),
        'MediumTurquoise': RGBColor(72, 209, 204),
        'MediumVioletRed': RGBColor(199, 21, 133),
        'MidnightBlue': RGBColor(25, 25, 112),
        'MintCream': RGBColor(245, 255, 250),
        'MistyRose': RGBColor(255, 228, 225),
        'Moccasin': RGBColor(255, 228, 181),
        'NavajoWhite': RGBColor(255, 222, 173),
        'Navy': RGBColor(0, 0, 128),
        'OldLace': RGBColor(253, 245, 230),
        'Olive': RGBColor(128, 128, 0),
        'OliveDrab': RGBColor(107, 142, 35),
        'Orange': RGBColor(255, 165, 0),
        'OrangeRed': RGBColor(255, 69, 0),
        'Orchid': RGBColor(218, 112, 214),
        'PaleGoldenRod': RGBColor(238, 232, 170),
        'PaleGreen': RGBColor(152, 251, 152),
        'PaleTurquoise': RGBColor(175, 238, 238),
        'PaleVioletRed': RGBColor(219, 112, 147),
        'PapayaWhip': RGBColor(255, 239, 213),
        'PeachPuff': RGBColor(255, 218, 185),
        'Peru': RGBColor(205, 133, 63),
        'Pink': RGBColor(255, 192, 203),
        'Plum': RGBColor(221, 160, 221),
        'PowderBlue': RGBColor(176, 224, 230),
        'Purple': RGBColor(128, 0, 128),
        'Red': RGBColor(255, 0, 0),
        'RosyBrown': RGBColor(188, 143, 143),
        'RoyalBlue': RGBColor(65, 105, 225),
        'SaddleBrown': RGBColor(139, 69, 19),
        'Salmon': RGBColor(250, 128, 114),
        'SandyBrown': RGBColor(244, 164, 96),
        'SeaGreen': RGBColor(46, 139, 87),
        'SeaShell': RGBColor(255, 245, 238),
        'Sienna': RGBColor(160, 82, 45),
        'Silver': RGBColor(192, 192, 192),
        'SkyBlue': RGBColor(135, 206, 235),
        'SlateBlue': RGBColor(106, 90, 205),
        'SlateGray': RGBColor(112, 128, 144),
        'Snow': RGBColor(255, 250, 250),
        'SpringGreen': RGBColor(0, 255, 127),
        'SteelBlue': RGBColor(70, 130, 180),
        'Tan': RGBColor(210, 180, 140),
        'Teal': RGBColor(0, 128, 128),
        'Thistle': RGBColor(216, 191, 216),
        'Tomato': RGBColor(255, 99, 71),
        'Turquoise': RGBColor(64, 224, 208),
        'Violet': RGBColor(238, 130, 238),
        'Wheat': RGBColor(245, 222, 179),
        'White': RGBColor(255, 255, 255),
        'WhiteSmoke': RGBColor(245, 245, 245),
        'Yellow': RGBColor(255, 255, 0),
        'YellowGreen': RGBColor(154, 205, 50),
        }

        if color_str in color_names:
            return color_names[color_str]
        
        # 处理带#前缀的十六进制颜色
        if color_str.startswith('#'):
            hex_str = color_str[1:]
            
            # 支持3位简写格式
            if len(hex_str) == 3:
                try:
                    r = int(hex_str[0]*2, 16)
                    g = int(hex_str[1]*2, 16)
                    b = int(hex_str[2]*2, 16)
                    return RGBColor(r, g, b)
                except ValueError:
                    return None
            
            # 处理6位格式
            if len(hex_str) == 6:
                try:
                    r = int(hex_str[0:2], 16)
                    g = int(hex_str[2:4], 16)
                    b = int(hex_str[4:6], 16)
                    return RGBColor(r, g, b)
                except ValueError:
                    return None
        
        # 支持rgb/rgba格式
        if color_str.startswith('rgb('):
            match = re.match(r'rgb\((\d+),\s*(\d+),\s*(\d+)\)', color_str)
            if match:
                return RGBColor(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        
        if color_str.startswith('rgba('):
            match = re.match(r'rgba\((\d+),\s*(\d+),\s*(\d+),\s*[\d.]+\)', color_str)
            if match:
                return RGBColor(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        
        return None

    def _apply_run_style(self, run: Run, style_state: Dict[str, Any]):
        """将样式状态应用到Run对象 - 优化颜色处理"""
        # 字体样式
        run.bold = style_state.get('bold', False)
        run.italic = style_state.get('italic', False)
        run.underline = style_state.get('underline', False)
        
        # 字体大小
        font_size = style_state.get('font_size')
        if font_size:
            run.font.size = font_size
        
        # 字体颜色 - 只在有设置时应用
        color = style_state.get('color')
        if color:
            run.font.color.rgb = color
        
        # 应用高亮颜色（文字高亮）
        highlight_color = style_state.get('highlight_color')
        if highlight_color:
            run.font.highlight_color = highlight_color
        
        # 应用字体
        font_family = style_state.get('font_family')
        if font_family:
            run.font.name = font_family
            # 设置中文字体
            rPr = run._element.get_or_add_rPr()
            rPr.rFonts.set(qn('w:eastAsia'), font_family)
        else:
            # 应用默认字体（不设置颜色）
            self.apply_default_font(run)

    def set_paragraph_background(self, paragraph: Paragraph, bg_color: RGBColor):
        """设置段落背景色（底纹）"""
        # 创建底纹元素
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color.rgb_hex}"/>')
        # 获取段落属性
        pPr = paragraph._element.get_or_add_pPr()
        # 添加底纹
        pPr.append(shd)

    def apply_default_font(self, run: Run):
        """应用默认字体（不设置颜色）"""
        run.font.name = DocxFontEnum.TIMES_NEW_ROMAN
        rPr = run._element.get_or_add_rPr()
        rPr.rFonts.set(qn('w:eastAsia'), DocxFontEnum.SONG_TI)

    def apply_paragraph_styles(self, paragraph: Paragraph, style_str: str):
        """将 HTML 段落的 style 应用于 docx 段落"""
        styles = self._parse_style_string(style_str)

        # 行间距
        if "line-height" in styles:
            line_height = styles["line-height"].strip()
            try:
                if line_height.replace('.', '', 1).isdigit():
                    paragraph.paragraph_format.line_spacing = float(line_height)
            except ValueError:
                pass

        # 段前段后间距
        if "margin-top" in styles or "margin-bottom" in styles:
            if "margin-top" in styles:
                mt = styles["margin-top"].strip()
                if mt.endswith("pt"):
                    try:
                        paragraph.paragraph_format.space_before = Pt(float(mt[:-2]))
                    except ValueError:
                        pass
            if "margin-bottom" in styles:
                mb = styles["margin-bottom"].strip()
                if mb.endswith("pt"):
                    try:
                        paragraph.paragraph_format.space_after = Pt(float(mb[:-2]))
                    except ValueError:
                        pass

        # 文本对齐
        if "text-align" in styles:
            alignment = self.map_text_align_to_docx(styles["text-align"])
            paragraph.alignment = alignment

    def handle_table(self, doc: Document, table_tag: Tag, parent_style: Dict[str, Any]):
        rows = table_tag.find_all('tr')
        if not rows:
            return

        # 计算最大列数
        num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        table = doc.add_table(rows=len(rows), cols=num_cols)
        
        # 应用表格样式
        if table_tag.has_attr('style'):
            self.apply_table_styles(table, table_tag['style'])
        else:
            table.style = 'Normal Table'

        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j >= num_cols:
                    continue

                doc_cell = table.cell(i, j)
                # 清除单元格默认段落
                for p in doc_cell.paragraphs:
                    p.clear()

                # 应用单元格样式
                if cell.has_attr('style'):
                    self.apply_cell_styles(doc_cell, cell['style'])
                
                # 创建段落
                paragraph = doc_cell.add_paragraph()
                self.set_paragraph_alignment(paragraph, cell)
                
                # 创建样式状态
                style_state = parent_style.copy()
                if cell.has_attr('style'):
                    self._update_style_from_attributes(style_state, cell)
                
                # 处理单元格内容
                self._process_children(doc, paragraph, cell, style_state)

                # 表头加粗
                if cell.name == 'th':
                    for run in paragraph.runs:
                        run.bold = True

    def apply_table_styles(self, table: Table, style_str: str):
        """应用表格样式"""
        styles = self._parse_style_string(style_str)
        
        # 设置表格宽度
        if "width" in styles:
            width = styles["width"].strip()
            if width.endswith("%"):
                try:
                    percent = float(width[:-1]) / 100.0
                    table.autofit = True
                    #table.width = Inches(8 * percent)
                except ValueError:
                    pass
        
        # 设置边框合并
        if "border-collapse" in styles and styles["border-collapse"] == "collapse":
            table.style = "Normal Table"

    def apply_cell_styles(self, cell: _Cell, style_str: str):
        """应用单元格样式"""
        styles = self._parse_style_string(style_str)
        
        # 设置单元格宽度
        if "width" in styles:
            width = styles["width"].strip()
            if width.endswith("%"):
                try:
                    percent = float(width[:-1]) / 100.0
                    cell.width = Inches(8 * percent) 
                except ValueError:
                    pass
        
        # 设置单元格背景色
        if "background-color" in styles:
            bg_color = styles["background-color"].strip()
            parsed_bg_color = self._parse_color(bg_color)
            if parsed_bg_color:
                # 创建底纹元素
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{parsed_bg_color.rgb_hex}"/>')
                # 获取单元格属性
                tcPr = cell._tc.get_or_add_tcPr()
                # 添加底纹
                tcPr.append(shd)

    def handle_list(self, doc: Document, list_tag: Tag, parent_style: Dict[str, Any]):
        """处理 <ul> 或 <ol> 列表标签，支持嵌套和样式"""
        list_style = 'List Bullet' if list_tag.name == 'ul' else 'List Number'
        
        # 根据嵌套层级调整缩进
        indent_level = min(self._list_level, 5)  # 最大支持5级缩进

        for item in list_tag.find_all('li', recursive=False):
            paragraph = doc.add_paragraph(style=list_style)
            self._current_paragraph = paragraph

            # 设置缩进
            if indent_level > 1:
                paragraph.paragraph_format.left_indent = Inches(0.25 * indent_level)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)

            # 设置默认段落间距
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)

            # 解析并应用 <li> 或 <ul> 的 style 样式
            if item.has_attr('style'):
                self.apply_paragraph_styles(paragraph, item['style'])
            elif list_tag.has_attr('style'):
                self.apply_paragraph_styles(paragraph, list_tag['style'])
            
            # 创建样式状态
            style_state = parent_style.copy()
            if item.has_attr('style'):
                self._update_style_from_attributes(style_state, item)
            
            # 解析内联内容
            self._process_children(doc, paragraph, item, style_state)
            self._current_paragraph = None
            
            # 应用块级背景色
            if style_state.get('background_color'):
                self.set_paragraph_background(paragraph, style_state['background_color'])

            # 处理嵌套列表
            nested_lists = item.find_all(['ul', 'ol'], recursive=False)
            for nested_list in nested_lists:
                self.handle_list(doc, nested_list, style_state)

    def set_paragraph_alignment(self, paragraph: Paragraph, tag: Tag):
        """设置段落对齐方式，支持align属性和style属性"""
        # 1. 检查align属性
        if tag.has_attr('align'):
            alignment = self.map_text_align_to_docx(tag['align'])
            paragraph.alignment = alignment
            return
        
        # 2. 检查style属性中的text-align
        if tag.has_attr('style'):
            styles = self._parse_style_string(tag['style'])
            if 'text-align' in styles:
                alignment = self.map_text_align_to_docx(styles['text-align'])
                paragraph.alignment = alignment
                return
        
        # 3. 检查父元素的样式
        parent = tag.parent
        while parent:
            if parent.has_attr('align'):
                alignment = self.map_text_align_to_docx(parent['align'])
                paragraph.alignment = alignment
                return
            if parent.has_attr('style'):
                styles = self._parse_style_string(parent['style'])
                if 'text-align' in styles:
                    alignment = self.map_text_align_to_docx(styles['text-align'])
                    paragraph.alignment = alignment
                    return
            parent = parent.parent

    @staticmethod
    def get_text_align_from_style(style_str):
        styles = dict([s.split(":", 1) for s in style_str.split(";") if ":" in s])
        return styles.get("text-align")

    @staticmethod
    def map_text_align_to_docx(text_align_value):
        if not text_align_value:
            return WD_ALIGN_PARAGRAPH.LEFT

        text_align_value = text_align_value.lower().strip()
        align_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        return align_map.get(text_align_value, WD_ALIGN_PARAGRAPH.LEFT)

    def _is_processed(self, tag: Tag) -> bool:
        return id(tag) in self._processed_tags

    def _mark_as_processed(self, tag: Tag):
        self._processed_tags.add(id(tag))
